import io

def load_contract(uploaded_file) -> str:
    """Extract text from Streamlit UploadedFile. Supports PDF, DOCX, TXT."""
    filename = uploaded_file.name.lower()
    file_bytes = uploaded_file.read()

    if filename.endswith(".pdf"):
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                pages = [p.extract_text() for p in pdf.pages if p.extract_text()]
            text = "\n\n".join(pages)
            if not text.strip():
                raise ValueError("PDF is empty or image-only. Use a text-based PDF.")
            return text
        except ImportError:
            try:
                import PyPDF2
                reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                pages = [p.extract_text() for p in reader.pages if p.extract_text()]
                return "\n\n".join(pages)
            except ImportError:
                raise ValueError("Install pdfplumber: pip install pdfplumber")

    elif filename.endswith(".docx"):
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            paras = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip(): paras.append(cell.text.strip())
            return "\n\n".join(paras)
        except ImportError:
            raise ValueError("Install python-docx: pip install python-docx")

    elif filename.endswith(".txt"):
        try:    return file_bytes.decode("utf-8")
        except: return file_bytes.decode("latin-1")

    else:
        raise ValueError(f"Unsupported format: {uploaded_file.name}. Use PDF, DOCX, or TXT.")

def get_word_count(text: str) -> int:
    return len(text.split()) if text else 0